import json
from io import BytesIO
from pathlib import Path
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from sqlalchemy import desc, func, or_
from sqlalchemy.orm import Session

from app import models, schemas
from app.database import get_db
from app.features.auth.dependencies import get_current_user

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None


router = APIRouter()


DEFAULT_RAG_DOCS = [
    {
        "title": "高血压居家管理指南",
        "category": "慢性病管理",
        "content": "高血压患者建议坚持每日晨起与睡前监测血压，记录收缩压、舒张压和心率变化。饮食方面应减少高盐、腌制和高脂食物摄入，保持规律作息与中等强度运动。若连续多日血压异常，应及时就医并根据医生建议调整用药。",
        "source": "系统示例数据",
        "tags": ["高血压", "居家监测", "慢病"],
        "is_active": True,
    },
    {
        "title": "糖尿病饮食搭配建议",
        "category": "饮食营养",
        "content": "糖尿病人群应优先选择低升糖指数食物，如全谷物、深色蔬菜、豆类和优质蛋白。三餐定时定量，避免暴饮暴食，同时关注主食摄入比例与加餐安排。建议配合血糖监测，逐步找到适合自己的饮食节奏。",
        "source": "系统示例数据",
        "tags": ["糖尿病", "营养", "控糖"],
        "is_active": True,
    },
    {
        "title": "焦虑情绪自我调节方法",
        "category": "心理健康",
        "content": "当出现持续紧张、烦躁或睡眠不佳时，可尝试深呼吸训练、正念冥想和规律运动。每天安排固定的放松时间，减少长期高压状态。若情绪问题持续影响生活和工作，应主动寻求专业心理支持。",
        "source": "系统示例数据",
        "tags": ["焦虑", "放松", "睡眠"],
        "is_active": True,
    },
    {
        "title": "春季传染病预防要点",
        "category": "疾病预防",
        "content": "春季气温变化大，建议保持室内通风，注意勤洗手、戴口罩和规律作息。出现发热、咳嗽等症状时应减少聚集活动，并及时就医。学校和家庭应加强日常环境清洁，降低交叉感染风险。",
        "source": "系统示例数据",
        "tags": ["预防", "传染病", "春季"],
        "is_active": True,
    },
    {
        "title": "办公室久坐人群运动建议",
        "category": "运动康复",
        "content": "久坐人群每工作 45 到 60 分钟应起身活动 3 到 5 分钟，可进行肩颈拉伸、腰背舒展和快走。每周建议安排至少 150 分钟中等强度运动，并关注核心力量训练，以减轻腰背不适和代谢风险。",
        "source": "系统示例数据",
        "tags": ["久坐", "拉伸", "康复"],
        "is_active": True,
    },
]


def _parse_tags(raw_tags: Optional[str]) -> List[str]:
    if not raw_tags:
        return []

    try:
        parsed = json.loads(raw_tags)
        if isinstance(parsed, list):
            return [str(tag) for tag in parsed if str(tag).strip()]
    except (json.JSONDecodeError, TypeError):
        pass

    return [tag.strip() for tag in raw_tags.split(",") if tag.strip()]


def _join_tags(tags: Optional[List[str]]) -> str:
    return json.dumps(tags or [], ensure_ascii=False)


def _to_rag_doc_response(doc: models.RagKnowledgeDocument) -> schemas.RagKnowledgeDocResponse:
    return schemas.RagKnowledgeDocResponse(
        id=doc.id,
        title=doc.title,
        category=doc.category,
        content=doc.content,
        source=doc.source,
        tags=_parse_tags(doc.tags),
        is_active=doc.is_active,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
    )


def _normalize_multiline_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    compact_lines = [line for line in lines if line]
    return "\n".join(compact_lines).strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        raise HTTPException(status_code=500, detail="缺少 PDF 解析依赖，请安装 pypdf")

    reader = PdfReader(BytesIO(file_bytes))
    contents: List[str] = []
    for page in reader.pages:
        contents.append(page.extract_text() or "")
    return _normalize_multiline_text("\n".join(contents))


def _extract_docx_text(file_bytes: bytes) -> str:
    if DocxDocument is None:
        raise HTTPException(status_code=500, detail="缺少 Word 解析依赖，请安装 python-docx")

    document = DocxDocument(BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
    return _normalize_multiline_text("\n".join(parts))


def _extract_import_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return _extract_docx_text(file_bytes)
    if suffix == ".doc":
        raise HTTPException(status_code=400, detail="暂不支持 .doc，请先转换为 .docx 或 PDF")
    raise HTTPException(status_code=400, detail="仅支持导入 PDF 或 DOCX 文件")


def _build_import_doc(
    *,
    title: str,
    category: str,
    content: str,
    source: Optional[str],
    tags: List[str],
    is_active: bool,
) -> models.RagKnowledgeDocument:
    return models.RagKnowledgeDocument(
        title=title,
        category=category,
        content=content,
        source=source,
        tags=_join_tags(tags),
        is_active=is_active,
    )


def _favorite_count_map(db: Session, article_ids: List[int]) -> dict[int, int]:
    if not article_ids:
        return {}

    rows = (
        db.query(models.ArticleFavorite.article_id, func.count(models.ArticleFavorite.id))
        .filter(models.ArticleFavorite.article_id.in_(article_ids))
        .group_by(models.ArticleFavorite.article_id)
        .all()
    )
    return {article_id: count for article_id, count in rows}


def _view_count_map(db: Session, article_ids: List[int]) -> dict[int, int]:
    if not article_ids:
        return {}

    rows = (
        db.query(models.ArticleReadHistory.article_id, func.coalesce(func.sum(models.ArticleReadHistory.read_count), 0))
        .filter(models.ArticleReadHistory.article_id.in_(article_ids))
        .group_by(models.ArticleReadHistory.article_id)
        .all()
    )
    return {article_id: int(count or 0) for article_id, count in rows}


def _favorite_status_map(db: Session, user_id: int, article_ids: List[int]) -> dict[int, bool]:
    if not article_ids:
        return {}

    rows = (
        db.query(models.ArticleFavorite.article_id)
        .filter(
            models.ArticleFavorite.user_id == user_id,
            models.ArticleFavorite.article_id.in_(article_ids),
        )
        .all()
    )
    return {article_id: True for article_id, in rows}


def _to_article_response(
    article: models.HealthArticle,
    favorite_count: int = 0,
    view_count: Optional[int] = None,
    is_favorited: bool = False,
) -> schemas.HealthArticleResponse:
    return schemas.HealthArticleResponse(
        id=article.id,
        title=article.title,
        category=article.category,
        summary=article.summary,
        content=article.content,
        cover_image=article.cover_image,
        tags=_parse_tags(article.tags),
        view_count=article.view_count if view_count is None else view_count,
        favorite_count=favorite_count,
        is_favorited=is_favorited,
        created_at=article.created_at,
        updated_at=article.updated_at,
    )


def _assert_category(category: str) -> None:
    if category not in schemas.ARTICLE_CATEGORIES:
        raise HTTPException(status_code=400, detail="无效的文章分类")


def _assert_admin(user: models.User) -> None:
    if user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="仅管理员可操作")


def _record_read_history(db: Session, user_id: int, article_id: int) -> None:
    history = (
        db.query(models.ArticleReadHistory)
        .filter(
            models.ArticleReadHistory.user_id == user_id,
            models.ArticleReadHistory.article_id == article_id,
        )
        .first()
    )

    if history:
        history.read_count += 1
        history.last_read_at = datetime.utcnow()
    else:
        db.add(
            models.ArticleReadHistory(
                user_id=user_id,
                article_id=article_id,
                read_count=1,
                last_read_at=datetime.utcnow(),
            )
        )


@router.get("/articles", response_model=schemas.HealthArticleListResponse)
async def list_articles(
    page: int = Query(1, ge=1),
    page_size: int = Query(12, ge=1, le=100),
    category: Optional[str] = None,
    keyword: Optional[str] = None,
    sort_by: str = Query("latest", pattern="^(latest|hot)$"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    query = db.query(models.HealthArticle)

    if category:
        query = query.filter(models.HealthArticle.category == category)

    if keyword:
        like = f"%{keyword}%"
        query = query.filter(
            or_(
                models.HealthArticle.title.ilike(like),
                models.HealthArticle.summary.ilike(like),
                models.HealthArticle.content.ilike(like),
            )
        )

    total = query.count()

    if sort_by == "hot":
        query = query.order_by(desc(models.HealthArticle.view_count), desc(models.HealthArticle.created_at))
    else:
        query = query.order_by(desc(models.HealthArticle.created_at))

    articles = query.offset((page - 1) * page_size).limit(page_size).all()
    article_ids = [article.id for article in articles]
    count_map = _favorite_count_map(db, article_ids)
    view_map = _view_count_map(db, article_ids)
    favorite_status_map = _favorite_status_map(db, current_user.id, article_ids)

    return schemas.HealthArticleListResponse(
        items=[
            _to_article_response(
                article,
                count_map.get(article.id, 0),
                view_map.get(article.id, article.view_count),
                favorite_status_map.get(article.id, False),
            )
            for article in articles
        ],
        total=total,
        page=page,
        page_size=page_size,
    )


@router.get("/articles/{article_id}", response_model=schemas.HealthArticleResponse)
async def get_article_detail(
    article_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    article = db.query(models.HealthArticle).filter(models.HealthArticle.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    article.view_count += 1
    _record_read_history(db, current_user.id, article_id)
    db.commit()
    db.refresh(article)

    count_map = _favorite_count_map(db, [article.id])
    view_map = _view_count_map(db, [article.id])
    favorite_status_map = _favorite_status_map(db, current_user.id, [article.id])
    return _to_article_response(
        article,
        count_map.get(article.id, 0),
        view_map.get(article.id, article.view_count),
        favorite_status_map.get(article.id, False),
    )


@router.post("/articles/{article_id}/favorite", response_model=schemas.FavoriteResponse)
async def favorite_article(
    article_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    article = db.query(models.HealthArticle).filter(models.HealthArticle.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    existing = (
        db.query(models.ArticleFavorite)
        .filter(
            models.ArticleFavorite.user_id == current_user.id,
            models.ArticleFavorite.article_id == article_id,
        )
        .first()
    )
    if not existing:
        db.add(models.ArticleFavorite(user_id=current_user.id, article_id=article_id))
        db.commit()

    count_map = _favorite_count_map(db, [article_id])
    return schemas.FavoriteResponse(
        article_id=article_id,
        is_favorited=True,
        favorite_count=count_map.get(article_id, 0),
    )


@router.delete("/articles/{article_id}/favorite", response_model=schemas.FavoriteResponse)
async def unfavorite_article(
    article_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    favorite = (
        db.query(models.ArticleFavorite)
        .filter(
            models.ArticleFavorite.user_id == current_user.id,
            models.ArticleFavorite.article_id == article_id,
        )
        .first()
    )

    if favorite:
        db.delete(favorite)
        db.commit()

    count_map = _favorite_count_map(db, [article_id])
    return schemas.FavoriteResponse(
        article_id=article_id,
        is_favorited=False,
        favorite_count=count_map.get(article_id, 0),
    )


@router.get("/favorites", response_model=schemas.HealthArticleListResponse)
async def list_favorites(
    page: int = Query(1, ge=1),
    page_size: int = Query(12, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    favorites_query = db.query(models.ArticleFavorite).filter(models.ArticleFavorite.user_id == current_user.id)
    total = favorites_query.count()

    favorites = (
        favorites_query.order_by(desc(models.ArticleFavorite.created_at))
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )

    articles = [favorite.article for favorite in favorites if favorite.article]
    article_ids = [article.id for article in articles]
    count_map = _favorite_count_map(db, article_ids)
    view_map = _view_count_map(db, article_ids)
    favorite_status_map = _favorite_status_map(db, current_user.id, article_ids)

    return schemas.HealthArticleListResponse(
        items=[
            _to_article_response(
                article,
                count_map.get(article.id, 0),
                view_map.get(article.id, article.view_count),
                favorite_status_map.get(article.id, False),
            )
            for article in articles
        ],
        total=total,
        page=page,
        page_size=page_size,
    )


@router.get("/read-history", response_model=List[schemas.ReadingHistoryResponse])
async def list_read_history(
    limit: int = Query(30, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    rows = (
        db.query(models.ArticleReadHistory)
        .filter(models.ArticleReadHistory.user_id == current_user.id)
        .order_by(desc(models.ArticleReadHistory.last_read_at))
        .limit(limit)
        .all()
    )

    return [
        schemas.ReadingHistoryResponse(
            article_id=row.article_id,
            article_title=row.article.title if row.article else "",
            category=row.article.category if row.article else "",
            last_read_at=row.last_read_at,
            read_count=row.read_count,
        )
        for row in rows
        if row.article
    ]


@router.get("/recommendations/home", response_model=schemas.HomepageRecommendationResponse)
async def get_home_recommendations(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    hot_articles = (
        db.query(models.HealthArticle)
        .order_by(desc(models.HealthArticle.view_count), desc(models.HealthArticle.created_at))
        .limit(6)
        .all()
    )
    latest_articles = db.query(models.HealthArticle).order_by(desc(models.HealthArticle.created_at)).limit(6).all()

    ids = list({article.id for article in hot_articles + latest_articles})
    count_map = _favorite_count_map(db, ids)
    view_map = _view_count_map(db, ids)
    favorite_status_map = _favorite_status_map(db, current_user.id, ids)

    return schemas.HomepageRecommendationResponse(
        hot_articles=[
            _to_article_response(
                article,
                count_map.get(article.id, 0),
                view_map.get(article.id, article.view_count),
                favorite_status_map.get(article.id, False),
            )
            for article in hot_articles
        ],
        latest_articles=[
            _to_article_response(
                article,
                count_map.get(article.id, 0),
                view_map.get(article.id, article.view_count),
                favorite_status_map.get(article.id, False),
            )
            for article in latest_articles
        ],
    )


@router.post("/admin/articles", response_model=schemas.HealthArticleResponse)
async def create_article(
    payload: schemas.HealthArticleCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)
    _assert_category(payload.category)

    article = models.HealthArticle(
        title=payload.title,
        category=payload.category,
        summary=payload.summary,
        content=payload.content,
        cover_image=payload.cover_image,
        tags=_join_tags(payload.tags),
    )
    db.add(article)
    db.commit()
    db.refresh(article)

    return _to_article_response(article, 0, 0, False)


@router.put("/admin/articles/{article_id}", response_model=schemas.HealthArticleResponse)
async def update_article(
    article_id: int,
    payload: schemas.HealthArticleUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)

    article = db.query(models.HealthArticle).filter(models.HealthArticle.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    update_data = payload.model_dump(exclude_unset=True)
    if "category" in update_data and update_data["category"]:
        _assert_category(update_data["category"])

    if "tags" in update_data:
        update_data["tags"] = _join_tags(update_data["tags"]) if update_data["tags"] is not None else _join_tags([])

    for key, value in update_data.items():
        setattr(article, key, value)

    db.commit()
    db.refresh(article)

    count_map = _favorite_count_map(db, [article.id])
    view_map = _view_count_map(db, [article.id])
    favorite_status_map = _favorite_status_map(db, current_user.id, [article.id])
    return _to_article_response(
        article,
        count_map.get(article.id, 0),
        view_map.get(article.id, article.view_count),
        favorite_status_map.get(article.id, False),
    )


@router.delete("/admin/articles/{article_id}")
async def delete_article(
    article_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)

    article = db.query(models.HealthArticle).filter(models.HealthArticle.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    db.delete(article)
    db.commit()
    return {"message": "文章已删除"}


@router.get("/admin/rag-docs", response_model=schemas.RagKnowledgeDocListResponse)
async def list_rag_docs(
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=100),
    keyword: Optional[str] = None,
    category: Optional[str] = None,
    active_only: bool = Query(False),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)
    query = db.query(models.RagKnowledgeDocument)

    if keyword:
        like = f"%{keyword}%"
        query = query.filter(
            or_(
                models.RagKnowledgeDocument.title.ilike(like),
                models.RagKnowledgeDocument.content.ilike(like),
                models.RagKnowledgeDocument.source.ilike(like),
            )
        )
    if category:
        query = query.filter(models.RagKnowledgeDocument.category == category)
    if active_only:
        query = query.filter(models.RagKnowledgeDocument.is_active.is_(True))

    total = query.count()
    items = (
        query.order_by(desc(models.RagKnowledgeDocument.updated_at))
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )

    return schemas.RagKnowledgeDocListResponse(
        items=[_to_rag_doc_response(item) for item in items],
        total=total,
        page=page,
        page_size=page_size,
    )


@router.post("/admin/rag-docs", response_model=schemas.RagKnowledgeDocResponse)
async def create_rag_doc(
    payload: schemas.RagKnowledgeDocCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)
    doc = models.RagKnowledgeDocument(
        title=payload.title,
        category=payload.category,
        content=payload.content,
        source=payload.source,
        tags=_join_tags(payload.tags),
        is_active=payload.is_active,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _to_rag_doc_response(doc)


@router.post("/admin/rag-docs/import", response_model=schemas.RagKnowledgeImportResponse)
async def import_rag_docs(
    files: List[UploadFile] = File(...),
    category: str = Form(...),
    source: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    is_active: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)

    if not files:
        raise HTTPException(status_code=400, detail="请至少选择一个文件")

    parsed_tags = _parse_tags(tags)
    imported_docs: List[models.RagKnowledgeDocument] = []
    skipped_files: List[str] = []

    for upload in files:
        filename = upload.filename or "未命名文件"
        file_bytes = await upload.read()

        if not file_bytes:
            skipped_files.append(f"{filename}（空文件）")
            continue

        try:
            content = _extract_import_text(filename, file_bytes)
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"文件 {filename} 解析失败：{exc}") from exc

        if not content:
            skipped_files.append(f"{filename}（未提取到有效文本）")
            continue

        title = Path(filename).stem.strip() or "导入知识文档"
        doc_source = source or filename
        imported_docs.append(
            _build_import_doc(
                title=title,
                category=category,
                content=content,
                source=doc_source,
                tags=parsed_tags,
                is_active=is_active,
            )
        )

    if not imported_docs:
        raise HTTPException(status_code=400, detail="没有可导入的有效文件")

    db.add_all(imported_docs)
    db.commit()
    for doc in imported_docs:
        db.refresh(doc)

    return schemas.RagKnowledgeImportResponse(
        items=[_to_rag_doc_response(doc) for doc in imported_docs],
        imported_count=len(imported_docs),
        skipped_files=skipped_files,
        message=f"成功导入 {len(imported_docs)} 个知识库文档",
    )


@router.post("/admin/rag-docs/seed-defaults", response_model=schemas.RagKnowledgeImportResponse)
async def seed_default_rag_docs(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)

    created_docs: List[models.RagKnowledgeDocument] = []
    skipped_files: List[str] = []

    for item in DEFAULT_RAG_DOCS:
        existing = (
            db.query(models.RagKnowledgeDocument)
            .filter(
                models.RagKnowledgeDocument.title == item["title"],
                models.RagKnowledgeDocument.category == item["category"],
            )
            .first()
        )
        if existing:
            skipped_files.append(item["title"])
            continue

        doc = _build_import_doc(
            title=item["title"],
            category=item["category"],
            content=item["content"],
            source=item["source"],
            tags=item["tags"],
            is_active=item["is_active"],
        )
        db.add(doc)
        created_docs.append(doc)

    db.commit()
    for doc in created_docs:
        db.refresh(doc)

    return schemas.RagKnowledgeImportResponse(
        items=[_to_rag_doc_response(doc) for doc in created_docs],
        imported_count=len(created_docs),
        skipped_files=skipped_files,
        message=(
            f"已导入 {len(created_docs)} 条示例知识"
            if created_docs
            else "示例知识已存在，未重复导入"
        ),
    )


@router.put("/admin/rag-docs/{doc_id}", response_model=schemas.RagKnowledgeDocResponse)
async def update_rag_doc(
    doc_id: int,
    payload: schemas.RagKnowledgeDocUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)
    doc = db.query(models.RagKnowledgeDocument).filter(models.RagKnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="知识库文档不存在")

    update_data = payload.model_dump(exclude_unset=True)
    if "tags" in update_data:
        update_data["tags"] = _join_tags(update_data["tags"] or [])

    for key, value in update_data.items():
        setattr(doc, key, value)

    db.commit()
    db.refresh(doc)
    return _to_rag_doc_response(doc)


@router.delete("/admin/rag-docs/{doc_id}")
async def delete_rag_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    _assert_admin(current_user)
    doc = db.query(models.RagKnowledgeDocument).filter(models.RagKnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="知识库文档不存在")

    db.delete(doc)
    db.commit()
    return {"message": "知识库文档已删除"}
